# -*- coding: utf-8 -*-
"""
Created on Tue Jun  9 14:10:44 2020

@author: pielsticker
"""
import os
import shelve
import numpy as np
import json
from matplotlib import pyplot as plt
from matplotlib.ticker import MaxNLocator


from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.shared import Pt


#%%                
                
class TrainingGraphs():
    def __init__(self, history, model_name, time):
        self.history = history
        root_dir = os.getcwd().partition('network')[0]
        
        self.fig_file_name = root_dir + '\\figures\\' +\
            time + '_' + model_name + '\\'
        
        self.plot_loss()
        self.plot_accuracy()
        
    def plot_loss(self):
        # summarize history for loss
        fig, ax = plt.subplots(figsize=(9,5))
        ax.plot(self.history['loss'], linewidth = 3)
        ax.plot(self.history['val_loss'], linewidth = 3)
        ax.set_title('Loss')
        ax.set_ylabel('Cross Entropy Loss')
        ax.set_xlabel('Epoch')
        ax.set_xticks(range(0, len(self.history['loss']), 1))
        ax.set_xticklabels(range(1, len(self.history['accuracy'])+1, 1))
        ax.legend(['Train', 'Validation'], loc='lower right')
        fig_name = self.fig_file_name + 'loss.png' 
        fig.savefig(fig_name)
        plt.show()
        
    def plot_accuracy(self):
        # summarize history for loss
        fig, ax = plt.subplots(figsize=(9,5))
        ax.plot(self.history['accuracy'], linewidth = 3)
        ax.plot(self.history['val_accuracy'], linewidth = 3)
        ax.set_title('Accuracy')
        ax.set_ylabel('Classification Accuracy')
        ax.set_xlabel('Epoch')
        ax.set_xticks(range(0, len(self.history['accuracy']), 1))
        ax.set_xticklabels(range(1, len(self.history['accuracy'])+1, 1))
        ax.legend(['Train', 'Validation'], loc='lower right')
        fig_name = self.fig_file_name + 'accuracy.png' 
        fig.savefig(fig_name)
        plt.show()
        
        

class Report:
    def __init__(self, dir_name = ''):
        self.document = Document()
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
           
        # Get the data
        root_dir = os.getcwd().partition('network')[0]
        self.model_dir = root_dir + '\\saved_models\\' + dir_name + '\\'
        self.log_dir = root_dir + '\\logs\\' + dir_name + '\\'
        self.fig_dir = root_dir + '\\figures\\' + dir_name + '\\'
        
        self.name_data, self.class_dist, self.train_data, self.model_summary =\
            self.get_hyperparams()
        self.results = self.get_results()
        
        self.filename = self.log_dir + 'results.docx'     
        self.create_document()
        
    def create_document(self):
        self.document.add_heading('Training report', 0)
        
        
        self.document.add_heading('Data:', 1)
        
        name_table = self.document.add_table(
            rows = len(self.name_data.keys()), cols = 2)
        for key, value in self.name_data.items():
            j = int(list(self.name_data.keys()).index(key))
            name_table.cell(j, 0).text = key + ':'
            name_table.cell(j, 1).text = str(value)
            
          
        self.document.add_heading('Class distribution:', 1)  
        
        dist_table = self.document.add_table(
            rows = len(self.class_dist.keys())+1,
            cols = len(next(iter(self.class_dist.values())))+1)
        
        dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, name in enumerate(self.name_data['Labels']):
            dist_table.cell(0, i+1).text = name
        for item, param in self.class_dist.items():
            j = int(list(self.class_dist.keys()).index(item))+1
            dist_table.cell(j, 0).text = item
            
            for key, value in self.class_dist[item].items():
                k = int(key)+1
                dist_table.cell(j, k).text = str(value)
        
        self.document.add_page_break()
        self.document.add_heading('Training parameters:', 1)
        
        train_table = self.document.add_table(
            rows = len(self.train_data.keys()),
            cols = 2)   
        
        for key, value in self.train_data.items():
            j = int(list(self.train_data.keys()).index(key))
            train_table.cell(j, 0).text = key + ':'
            train_table.cell(j, 1).text = str(value)
            
            
        self.document.add_heading('Model architecture', 1)
        
        
        par = self.document.add_paragraph(self.model_summary)
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.document.add_page_break()
        self.document.add_heading('Loss & accuracy', 1)
        
        loss_file = self.fig_dir + 'loss.png'
        acc_file = self.fig_dir + 'accuracy.png'
        self.document.add_picture(loss_file, width=Cm(12))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_picture(acc_file, width=Cm(12))  
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        self.document.add_heading('Results', 1)
        
        result_table = self.document.add_table(rows = 2, 
                                              cols = 2)
        result_table.cell(0, 0).text = 'Test loss:'
        result_table.cell(0, 1).text = str(np.round(
            self.results['test_loss'],decimals = 3))
        result_table.cell(1, 0).text = 'Test accuracy:'
        result_table.cell(1, 1).text = str(np.round(
            self.results['test_accuracy'],decimals = 3))
        
        for row in result_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER
                
        self.document.add_page_break()
        self.document.add_heading('Predictions for 5 random examples', 1)
        
        self.document.add_heading('Training data', 2)
        
        r = np.random.randint(0, self.results['y_train'].shape[0]-5)
        pred_train_5 = self.results['pred_train'][r:r+5,:]
        y_train_5 = self.results['y_train'][r:r+5,:] 
        
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = 'Predictions:'
        run.font.underline = True
        pred_train_table = self.add_result_table(pred_train_5)
        
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = 'Correct labels:'
        run.font.underline = True
        y_train_table = self.add_result_table(y_train_5)
           
        
        self.document.add_heading('Test data', 2)  
        
        s = np.random.randint(0, self.results['y_test'].shape[0]-5)
        pred_test_5 = self.results['pred_test'][s:s+5,:] 
        y_test_5 = self.results['y_test'][s:s+5,:]
        
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = 'Predictions:'
        run.font.underline = True
        pred_test_table = self.add_result_table(pred_train_5)

        
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = 'Correct labels:'
        run.font.underline = True
        y_test_table = self.add_result_table(y_train_5)
        
        
    def add_result_table(self, data_array):
        new_table = self.document.add_table(rows = data_array.shape[0]+1, 
                                           cols = data_array.shape[1]) 
        for row in new_table.rows:
            row.height = Cm(0.5)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
        for i, name in enumerate(self.name_data['Labels']):
            new_table.cell(0, i).text = name
            new_table.cell(0, i).paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER
            
        if data_array.dtype == 'float32':
            a = np.around(data_array, decimals = 4)
            row_sums = a.sum(axis = 1)
            data_array = a / row_sums[:, np.newaxis]
            data_array = np.around(data_array, decimals = 2)

        for i in range(data_array.shape[0]):
            for j in range(data_array.shape[1]):
                new_table.cell(i+1, j).text = str(data_array[i,j])
                new_table.cell(i+1, j).paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER
        

                    
    def get_hyperparams(self):
        hyperparam_file_name = self.model_dir + 'hyperparameters.json'
        with open(hyperparam_file_name) as json_file:
            data_dict = json.load(json_file)
            
        name_data = {
            'Name' : data_dict['model_name'],
            'Time created' : data_dict['datetime'], 
            'No. of classes' : data_dict['num_of_classes'],
            'Labels': data_dict['Labels'],
            'Total no. of samples' : data_dict['Total no. of samples'],
            'Train-test-split' : data_dict['train_test_split_percentage'],
            'Train-val-split' : data_dict['train_val_split_percentage'],
            'No. of training samples' : data_dict['No. of training samples'],
            'No. of validation samples' : data_dict['No. of validation samples'],
            'No. of test samples' : data_dict['No. of test samples'],
            'Shape of each sample' : data_dict['Shape of each sample']
            }
        
        class_distribution = data_dict['class_distribution']

        train_data = {            
            'Optimizer' : data_dict['optimizer'],
            'Learning rate' : data_dict['learning rate'],
            'Loss function' : data_dict['loss'],
            'Epochs trained' : data_dict['epochs_trained'],
            'Batch size' : data_dict['batch_size']}
        
        model_summary = data_dict['model_summary']

        
        return name_data, class_distribution, train_data, model_summary
        
            
    def get_results(self):
        data = {}
        file_name = self.model_dir + 'vars'
        with shelve.open(file_name) as shelf:
            for key in shelf:
                data[key] = shelf[key]
                
        return data
        
    
    def write(self):
        self.document.save(self.filename)
        print('Report saved!')
    

#%% 
if __name__ == "__main__":
    dir_name = '20200615_11h28m_Fe_single_4_classes_CNN_simple'
    rep = Report(dir_name)  
    data = rep.get_results()
    summary = rep.model_summary
    rep.write()  